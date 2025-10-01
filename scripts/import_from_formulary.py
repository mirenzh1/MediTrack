import os
import sys
from docx import Document
from supabase import create_client, Client
from dotenv import load_dotenv

# Fix Windows console encoding issues
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

# Load environment variables
load_dotenv('../.env')

# Get Supabase credentials
SUPABASE_URL = os.getenv('VITE_SUPABASE_URL')
SUPABASE_KEY = os.getenv('VITE_SUPABASE_ANON_KEY')

# Initialize Supabase client
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

print("=" * 70)
print("IMPORT MEDICATIONS FROM 2025 EMORY FORMULARY")
print("=" * 70)

# Step 1: Read the DOCX file and extract Table 1
print("\n[1/3] Reading 2025 Emory Formulary.docx...")
try:
    doc = Document('../data/2025 Emory Formulary.docx')

    if not doc.tables:
        print("   ✗ No tables found in document")
        exit(1)

    # Get Table 1 (first table)
    table = doc.tables[0]

    # Extract data from table
    medications = []
    for i, row in enumerate(table.rows):
        if i == 0:  # Skip header row
            continue

        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 3 and cells[0] and cells[0] != 'Name':
            name = cells[0]
            strength = cells[1]
            quantity = cells[2]

            medications.append({
                'name': name,
                'strength': strength,
                'quantity': quantity
            })

    print(f"   ✓ Extracted {len(medications)} medications from Table 1")

except Exception as e:
    print(f"   ✗ Error reading DOCX: {str(e)}")
    exit(1)

# Step 2: Delete all existing medications
print("\n[2/3] Deleting all existing medications from Supabase...")
try:
    # Get all medications
    all_meds = supabase.table('medications').select('id').execute()
    count = len(all_meds.data)

    if count > 0:
        print(f"   Found {count} existing medications")

        deleted = 0
        for med in all_meds.data:
            supabase.table('medications').delete().eq('id', med['id']).execute()
            deleted += 1
            if deleted % 10 == 0:
                print(f"   Deleted {deleted}/{count}...")

        print(f"   ✓ Deleted {deleted} medications")
    else:
        print("   ✓ No medications to delete")

except Exception as e:
    print(f"   ✗ Error deleting medications: {str(e)}")
    print("   Continuing with import...")

# Step 3: Import medications
print("\n[3/3] Importing medications from formulary...")

success_count = 0
error_count = 0
errors = []

def parse_quantity(qty_str):
    """Parse quantity string to get stock number"""
    qty_str = str(qty_str).strip().lower()

    # Handle special cases
    if 'dispense on-site' in qty_str or 'on-site' in qty_str:
        return 0  # On-site only, no stock to track

    if 'x' == qty_str:
        return 0

    # Extract first number from string
    import re
    match = re.search(r'(\d+)', qty_str)
    if match:
        return int(match.group(1))

    return 0

def determine_dosage_form(name, strength):
    """Determine dosage form from medication name and strength"""
    name_lower = name.lower()

    if 'inhaler' in name_lower:
        return 'inhaler'
    elif 'liquid' in name_lower or 'susp' in name_lower or 'syrup' in name_lower:
        return 'liquid'
    elif 'cream' in name_lower or 'ointment' in name_lower or 'gel' in name_lower:
        return 'topical'
    elif 'drop' in name_lower:
        return 'drops'
    elif 'spray' in name_lower:
        return 'spray'
    elif 'caps' in name_lower or 'capsule' in name_lower:
        return 'capsule'
    elif 'vitamin' in name_lower or 'rub' in name_lower:
        return 'tablet'
    else:
        return 'tablet'

for i, med in enumerate(medications, 1):
    try:
        name = med['name']
        strength = med['strength']
        quantity = med['quantity']

        # Parse stock quantity
        stock = parse_quantity(quantity)

        # Determine dosage form
        dosage_form = determine_dosage_form(name, strength)

        # Clean up medication name (remove dosage form info)
        clean_name = (name.replace(' caps', '')
                          .replace(' susp', '')
                          .replace(', chewable', '')
                          .replace(', liquid', '')
                          .replace(',  liquid', '')
                          .strip())

        # Prepare medication data
        medication_data = {
            'name': clean_name,
            'strength': strength if strength and strength != 'x' else None,
            'dosage_form': dosage_form,
            'is_active': True,
            'current_stock': stock,
            'notes': f"Formulary qty: {quantity}"
        }

        # Insert into Supabase
        result = supabase.table('medications').insert(medication_data).execute()

        success_count += 1
        print(f"   ✓ [{success_count:2d}] {clean_name:40s} {strength:15s} Stock: {stock:3d}")

    except Exception as e:
        error_count += 1
        error_msg = f"Row {i} ({name}): {str(e)}"
        errors.append(error_msg)
        print(f"   ✗ Error: {error_msg}")

# Summary
print("\n" + "=" * 70)
print("IMPORT SUMMARY")
print("=" * 70)
print(f"✓ Successfully imported: {success_count} medications")
if error_count > 0:
    print(f"✗ Errors: {error_count}")
    print("\nError details:")
    for error in errors[:5]:
        print(f"  - {error}")
    if len(errors) > 5:
        print(f"  ... and {len(errors) - 5} more errors")

print("\n" + "=" * 70)
print("Done! Check Supabase dashboard to verify the data.")
print("=" * 70)
